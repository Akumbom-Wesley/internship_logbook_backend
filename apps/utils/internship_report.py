from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from apps.internships.models import Internship
import os
from datetime import datetime
from django.conf import settings
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import textwrap


class InternshipReportGenerateView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, internship_id):
        try:
            internship = Internship.objects.select_related(
                'student__user', 'student__department__school',
                'company', 'supervisor__user',
                'logbook'
            ).prefetch_related(
                'logbook__weekly_logs__logbook_entries'
            ).get(id=internship_id, student__user=request.user)

            if internship.status != 'completed':
                return Response({"error": "Internship must be completed before report generation."},
                                status=status.HTTP_400_BAD_REQUEST)

            if hasattr(internship, 'report_generated') and internship.report_generated:
                return Response({"error": "Report has already been generated for this internship."},
                                status=status.HTTP_400_BAD_REQUEST)

            # Configure Gemini API
            genai.configure(api_key=settings.GEMINI_API_KEY)
            model = genai.GenerativeModel('gemini-2.0-flash')

            # Generate content with specific instructions to avoid placeholders and options
            prompt_template = """
            Generate a professional internship report in first person perspective based on the following details.
            Do not provide multiple options or placeholders - generate complete, finalized content.
            Remove all markdown formatting like asterisks (*) and only include the actual content.

            Student Name: {student_name}
            Registration Number: {matricule}
            Level: {level}
            Department: {department}
            School: {school}

            Company Name: {company}
            Supervisor Name: {supervisor}
            Internship Duration: {start_date} to {end_date}
            Job Description: {job_description}

            Weekly Logs:
            {weekly_logs}

            The report should follow this structure:
            1. Title Page
            2. Dedication (single paragraph)
            3. Acknowledgment (single paragraph)
            4. Executive Summary
            5. Chapter 1: Introduction
               1.1 Overview
               1.2 Objectives
               1.3 Company Presentation
            6. Chapter 2: Internship Activities (detailed by week)
            7. Chapter 3: Technical Details
            8. Chapter 4: Skills and Lessons
            9. Chapter 5: Conclusion and Recommendations
            10. References

            Write in complete sentences and paragraphs, not bullet points.
            """

            # Format the prompt with actual data
            prompt = prompt_template.format(
                student_name=internship.student.user.full_name,
                matricule=internship.student.matricule_num,
                level=internship.student.level,
                department=internship.student.department.name,
                school=internship.student.department.school.name,
                company=internship.company.name,
                supervisor=internship.supervisor.user.full_name,
                start_date=internship.start_date.strftime('%B %d, %Y'),
                end_date=internship.end_date.strftime('%B %d, %Y'),
                job_description=internship.job_description,
                weekly_logs=self._format_weekly_logs(internship)
            )

            # Generate each section with specific instructions
            dedication = model.generate_content(
                f"Write a single paragraph dedication for an internship report by {internship.student.user.full_name}. "
                "Do not provide options - write the actual dedication text only. "
                "Example format: 'I dedicate this report to...'"
            ).text.replace('*', '')

            acknowledgment = model.generate_content(
                f"Write a single paragraph acknowledgment for an internship report thanking {internship.supervisor.user.full_name} "
                f"and {internship.company.name}. Do not provide options - write the actual acknowledgment text only. "
                "Example format: 'I would like to thank...'"
            ).text.replace('*', '')

            executive_summary = model.generate_content(
                f"Write a professional executive summary (about 150 words) for an internship at {internship.company.name} "
                f"as a {internship.job_description}. Focus on key achievements and learning outcomes."
            ).text.replace('*', '')

            introduction = model.generate_content(
                f"Write a professional introduction chapter (about 300 words) for an internship report at {internship.company.name}. "
                "Include: 1.1 Overview of the internship, 1.2 Clear objectives, and 1.3 Presentation of the company. "
                "Write in complete paragraphs, not bullet points."
            ).text.replace('*', '')

            activities = model.generate_content(
                f"Write a detailed chapter about internship activities based on these weekly logs: {self._format_weekly_logs(internship)}. "
                "Organize by week with specific tasks and accomplishments. Write in complete paragraphs."
            ).text.replace('*', '')

            technical_details = model.generate_content(
                "Write a technical details chapter describing projects worked on during the internship. "
                "Include technologies used and technical challenges overcome. Write in complete paragraphs."
            ).text.replace('*', '')

            skills_learned = model.generate_content(
                "Write a skills acquired and lessons learned chapter for an internship report. "
                "Include both technical and soft skills. Write in complete paragraphs."
            ).text.replace('*', '')

            conclusion = model.generate_content(
                "Write a conclusion and recommendations chapter for an internship report. "
                "Include 5.1 Conclusion summarizing the experience and 5.2 Recommendations for both the company and future interns. "
                "Write in complete paragraphs."
            ).text.replace('*', '')

            # Create Word document
            document = self._create_word_document(
                internship=internship,
                dedication=dedication,
                acknowledgment=acknowledgment,
                executive_summary=executive_summary,
                introduction=introduction,
                activities=activities,
                technical_details=technical_details,
                skills_learned=skills_learned,
                conclusion=conclusion
            )

            # Save to media folder
            file_path = self._save_report_to_media(document, internship)

            # Mark internship as having report generated
            internship.report_generated = True
            internship.save()

            return Response({
                "message": "Report generated successfully",
                "file_path": file_path
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            print(f"Error generating report: {str(e)}")
            return Response(
                {"error": f"Error generating report: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def _format_weekly_logs(self, internship):
        logs = []
        for week in internship.logbook.weekly_logs.all().order_by('week_no'):
            entries = [f"Week {week.week_no}:"]
            for entry in week.logbook_entries.filter(is_immutable=True).order_by('created_at'):
                entries.append(f"- {entry.created_at.strftime('%B %d, %Y')}: {entry.description}")
            logs.append("\n".join(entries))
        return "\n\n".join(logs)

    def _create_word_document(self, internship, **sections):
        document = Document()

        # Title Page
        self._add_title_page(document, internship)

        # Dedication
        document.add_page_break()
        self._add_section(document, "DEDICATION", sections['dedication'])

        # Acknowledgment
        document.add_page_break()
        self._add_section(document, "ACKNOWLEDGMENT", sections['acknowledgment'])

        # Executive Summary
        document.add_page_break()
        self._add_section(document, "EXECUTIVE SUMMARY", sections['executive_summary'])

        # Chapter 1: Introduction
        document.add_page_break()
        self._add_chapter(document, "CHAPTER 1: INTRODUCTION", sections['introduction'])

        # Chapter 2: Internship Activities
        document.add_page_break()
        self._add_chapter(document, "CHAPTER 2: INTERNSHIP ACTIVITIES", sections['activities'])

        # Chapter 3: Technical Details
        document.add_page_break()
        self._add_chapter(document, "CHAPTER 3: TECHNICAL DETAILS OF PROJECTS", sections['technical_details'])

        # Chapter 4: Skills Acquired
        document.add_page_break()
        self._add_chapter(document, "CHAPTER 4: SKILLS ACQUIRED AND LESSONS LEARNED", sections['skills_learned'])

        # Chapter 5: Conclusion
        document.add_page_break()
        self._add_chapter(document, "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS", sections['conclusion'])

        # References
        document.add_page_break()
        self._add_references(document)

        return document

    def _add_title_page(self, document, internship):
        # University Header
        title = document.add_heading('THE UNIVERSITY OF BAMENDA', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Department Info
        department = document.add_heading(f'DEPARTMENT OF {internship.student.department.name.upper()}', level=1)
        department.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report Title
        document.add_paragraph()
        report_title = document.add_heading('INTERNSHIP REPORT', level=0)
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Course Info
        document.add_paragraph()
        document.add_paragraph("COURSE TITLE: INDUSTRIAL ATTACHMENT", style='Heading 2')
        document.add_paragraph("COURSE CODE: [INSERT COURSE CODE]", style='Heading 2')

        # Student Info
        document.add_paragraph()
        document.add_paragraph("PRESENTED BY:", style='Heading 2')
        document.add_paragraph(internship.student.user.full_name, style='Heading 1')
        document.add_paragraph(f"REGISTRATION NUMBER: {internship.student.matricule_num}", style='Heading 2')
        document.add_paragraph(f"LEVEL: {internship.student.level}", style='Heading 2')

        # Instructor Info
        document.add_paragraph()
        document.add_paragraph("COURSE INSTRUCTOR: [INSERT COURSE INSTRUCTOR]", style='Heading 2')

    def _add_section(self, document, title, content):
        heading = document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Clean content and split into paragraphs
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        for p in paragraphs:
            document.add_paragraph(p)
            document.add_paragraph()  # Add space between paragraphs

    def _add_chapter(self, document, title, content):
        heading = document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Clean content and split into paragraphs
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        for p in paragraphs:
            if p.startswith('###'):
                # Subheading
                subheading = document.add_heading(p.replace('###', '').strip(), level=2)
            else:
                # Regular paragraph
                document.add_paragraph(p)
                document.add_paragraph()  # Add space between paragraphs

    def _add_references(self, document):
        heading = document.add_heading("REFERENCES", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        references = [
            "GitHub. (n.d.). GitHub documentation: Guide to version control, collaboration, and CI/CD.",
            "Google. (n.d.). Firebase documentation: Authentication and real-time database resources.",
            "JetBrains. (n.d.). Kotlin documentation: Comprehensive guide to Kotlin language.",
            "Spring.io. (n.d.). Spring Boot documentation: Build anything.",
            "PostgreSQL Global Development Group. (n.d.). PostgreSQL documentation."
        ]

        for ref in references:
            document.add_paragraph(ref, style='List Bullet')

    def _save_report_to_media(self, document, internship):
        try:
            # Create the directory path
            folder_name = 'internship_reports'
            media_folder_path = os.path.join(settings.MEDIA_ROOT, folder_name)

            # Ensure the directory exists
            os.makedirs(media_folder_path, exist_ok=True)

            # Create filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            student_name = internship.student.user.full_name.replace(' ', '_')
            filename = f"internship_report_{student_name}_{timestamp}.docx"
            file_path = os.path.join(media_folder_path, filename)

            # Save the document
            document.save(file_path)

            print(f"Report saved successfully at: {file_path}")

            return os.path.join(folder_name, filename)

        except Exception as e:
            print(f"Error saving report to media folder: {str(e)}")
            raise